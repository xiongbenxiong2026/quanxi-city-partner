# -*- coding: utf-8 -*-
"""
生成二部门九个岗位考核指标汇总文档
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, color_hex='4472C4', size=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(size))
        border.set(qn('w:color'), color_hex)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def set_run_font(run, font_name='微软雅黑', size=10.5, bold=False, color=None, italic=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_cover(doc):
    for _ in range(6):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run('全犀平台')
    set_run_font(run, size=28, bold=True)

    title_para2 = doc.add_paragraph()
    title_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = title_para2.add_run('二部门九岗位考核指标汇总')
    set_run_font(run2, size=28, bold=True)

    doc.add_paragraph()
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(sub_para.add_run('KPI Assessment Indicators Summary'), size=14, color=(128,128,128))

    doc.add_paragraph()
    doc.add_paragraph()
    for label, value in [('编制部门','全犀运营中心'),('生效日期','2026年5月19日'),('版本号','V1.0')]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(f'{label}：{value}'), size=12)

    doc.add_page_break()

def add_kpi_section(doc, section_title, positions):
    """添加一个大部门的考核指标章节"""
    # 章节标题
    heading = doc.add_heading(level=1)
    run = heading.add_run(section_title)
    set_run_font(run, size=16, bold=True, color=(0,0,0))
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for pos_code, pos_name, kpi_list in positions:
        # 岗位小标题
        h2 = doc.add_heading(level=2)
        run2 = h2.add_run(f'{pos_code} · {pos_name}')
        set_run_font(run2, size=13, bold=True, color=(0,0,0))

        # KPI表格
        table = doc.add_table(rows=len(kpi_list)+1, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ['序号', '考核指标', '目标值']
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            set_cell_shading(cell, '2E75B6')
            run_h = cell.paragraphs[0].add_run(h)
            set_run_font(run_h, bold=True, size=11, color=(255,255,255))
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for idx, (kpi_name, kpi_target) in enumerate(kpi_list, 1):
            row = table.rows[idx]
            row.cells[0].text = str(idx)
            row.cells[1].text = kpi_name
            row.cells[2].text = kpi_target
            for i, val in enumerate([str(idx), kpi_name, kpi_target]):
                set_run_font(row.cells[i].paragraphs[0].runs[0], size=10.5)
                row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 交替行底色
            if idx % 2 == 0:
                for cell in row.cells:
                    set_cell_shading(cell, 'D6E4F0')

        doc.add_paragraph()

    doc.add_page_break()

def main():
    doc = Document()

    # 默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(10.5)

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    add_cover(doc)

    # 市场部（招商中心）
    market_positions = [
        ('Ta', '流量招商', [
            ('月拓展流量方数量',   '≥20个/月'),
            ('流量投放落地率',     '≥80%（签约流量方30天内完成首次投放）'),
            ('流量上播率',         '≥60%（流量引至直播间的有效转化率）'),
            ('流量方满意度',       '≥85%好评率'),
        ]),
        ('La', '直播间招商', [
            ('月入驻商家数量',       '≥15家/月'),
            ('商家试播转化率',       '≥70%（签约商家7天内完成首播）'),
            ('商家品牌带货参与率',   '≥60%（签约商家30天内至少参与1次品牌带货）'),
            ('商家满意度',           '≥85%好评率'),
        ]),
        ('Ba', '品牌招商', [
            ('月入驻品牌方数量',       '≥20家/月'),
            ('品牌方商品上架率',       '≥80% SKU上架（入驻30天内）'),
            ('品牌直播场次完成率',     'A级≥10场/30天，B级≥6场/30天'),
            ('品牌方满意度',           '≥85%好评率'),
        ]),
    ]
    add_kpi_section(doc, '一、市场部（招商中心）考核指标', market_positions)

    # 运营部（运营中心）
    ops_positions = [
        ('Tb', '流量对接专员', [
            ('对接响应时效',     '交接后1个工作日内联系流量方'),
            ('粉丝绑定完成率',   '≥90%（签约流量方7天内完成粉丝绑定）'),
            ('首场投放完成率',   '≥85%（签约流量方30天内完成首次投放）'),
            ('流量方满意度',     '≥90%好评率'),
        ]),
        ('Tc', '流量运营专员', [
            ('流量方问题响应时效', '一般问题1个工作日内，重大问题2小时内'),
            ('流量方续约率',       '≥85%'),
            ('流量方月活跃率',     '≥80%（稳定流量方每月有粉丝购买记录）'),
            ('流量方满意度',       '≥88%好评率'),
            ('人均服务流量方数',   '20-50家/人'),
        ]),
        ('Lb', '直播对接专员', [
            ('对接响应时效',     '商家入驻交接后1个工作日内联系商家'),
            ('充值转化率',       '≥80%（对接商家3天内完成首次充值）'),
            ('首播完成率',       '≥85%（入驻后7天内完成首播）'),
            ('功能开通率',       '100%完成核心功能开通'),
            ('商家满意度',       '≥90%好评率'),
        ]),
        ('Lc', '直播运营专员', [
            ('商家问题响应时效',     '一般问题1个工作日内，重大问题2小时内'),
            ('商家续费率',           '≥85%'),
            ('品牌带货参与率',       '≥50%（稳定商家每月至少参与3次品牌带货）'),
            ('商家满意度',           '≥88%好评率'),
            ('人均服务商家数',       '30-50家/人'),
        ]),
        ('Bb', '品牌对接专员', [
            ('对接响应时效',       '交接后1个工作日内联系品牌方'),
            ('商品上架完成率',     '≥80% SKU上架（入驻30天内）'),
            ('首播完成率',         'A级≥6场/30天，B级≥3场/30天'),
            ('品牌方满意度',       '≥90%好评率'),
        ]),
        ('Bc', '品牌运营专员', [
            ('品牌方问题响应时效',   '一般问题1个工作日内，重大问题2小时内'),
            ('品牌直播场次完成率',   'A/S级≥6场/30天，B级≥3场/30天'),
            ('品牌方续约率',         '≥85%'),
            ('品牌方满意度',         '≥88%好评率'),
            ('人均服务品牌方数',     '50-100家/人'),
        ]),
    ]
    add_kpi_section(doc, '二、运营部（运营中心）考核指标', ops_positions)

    # 汇总总表
    heading = doc.add_heading(level=1)
    run = heading.add_run('三、九岗位考核指标总览表')
    set_run_font(run, size=16, bold=True)
    doc.add_paragraph()

    total_rows = 9 + 1  # header + 9 positions
    table = doc.add_table(rows=total_rows, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['部门', '代码', '岗位名称', '考核指标数', '核心指标摘要', '达标要求']
    header_colors = ['1F4E79', '2E75B6', '4472C4', '5B9BD5', '70AD47', 'ED7D31']
    for i, (h, c) in enumerate(zip(headers, header_colors)):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, c)
        run_h = cell.paragraphs[0].add_run(h)
        set_run_font(run_h, bold=True, size=10.5, color=(255,255,255))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    summary_data = [
        ('市场部', 'Ta', '流量招商',       '4项', '拓展量/落地率/上播率/满意度',  '月20个/落地率80%/满意度85%'),
        ('市场部', 'La', '直播间招商',     '4项', '入驻量/试播率/带货率/满意度',  '月15家/试播率70%/满意度85%'),
        ('市场部', 'Ba', '品牌招商',       '4项', '入驻量/上架率/直播场次/满意度', '月20家/上架率80%/满意度85%'),
        ('运营部', 'Tb', '流量对接专员',   '4项', '响应时效/绑定率/首投率/满意度', '1天响应/绑定率90%/满意度90%'),
        ('运营部', 'Tc', '流量运营专员',   '5项', '响应时效/续约率/活跃率/满意度', '续约率85%/活跃率80%/满意度88%'),
        ('运营部', 'Lb', '直播对接专员',   '5项', '响应时效/充值率/首播率/满意度', '1天响应/充值率80%/满意度90%'),
        ('运营部', 'Lc', '直播运营专员',   '5项', '响应时效/续费率/带货率/满意度', '续费率85%/带货率50%/满意度88%'),
        ('运营部', 'Bb', '品牌对接专员',   '4项', '响应时效/上架率/首播率/满意度', '1天响应/上架率80%/满意度90%'),
        ('运营部', 'Bc', '品牌运营专员',   '5项', '响应时效/直播场次/续约率/满意度','续约率85%/场次达标/满意度88%'),
    ]
    for idx, row_data in enumerate(summary_data, 1):
        row = table.rows[idx]
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = val
            set_run_font(cell.paragraphs[0].runs[0], size=10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if idx % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, 'D6E4F0')

    output_path = r'D:\全犀平台\运营管理\二部门九岗位考核指标汇总.docx'
    doc.save(output_path)
    print(f'文档已生成：{output_path}')

if __name__ == '__main__':
    main()
