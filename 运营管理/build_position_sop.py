# -*- coding: utf-8 -*-
"""
生成运营部岗位安排SOP文档
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

def set_run_font(run, font_name='微软雅黑', size=10.5, bold=False, color=None):
    """设置文字格式"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading_custom(doc, text, level=1):
    """添加自定义标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    set_run_font(run, size=(18 if level==1 else (14 if level==2 else 12)), bold=True, color=(0,0,0))
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_cover(doc):
    """添加封面"""
    # 空行占位
    for _ in range(6):
        doc.add_paragraph()

    # 标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('全犀平台运营中心')
    set_run_font(title_run, size=28, bold=True, color=(0, 0, 0))

    title_para2 = doc.add_paragraph()
    title_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run2 = title_para2.add_run('岗位安排SOP')
    set_run_font(title_run2, size=28, bold=True, color=(0, 0, 0))

    doc.add_paragraph()

    # 副标题
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run('Standard Operating Procedure')
    set_run_font(sub_run, size=14, color=(128, 128, 128))

    doc.add_paragraph()
    doc.add_paragraph()

    # 信息
    info_items = [
        ('编制部门', '全犀运营中心'),
        ('生效日期', '2026年5月19日'),
        ('版本号', 'V1.0'),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(f'{label}：{value}'), size=12)

    doc.add_page_break()

def add_toc_placeholder(doc):
    """添加目录占位提示"""
    heading = doc.add_heading('目  录', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        set_run_font(run, size=18, bold=True)

    toc_items = [
        '一、标准岗位代码对照表',
        '二、运营部岗位人员安排',
        '三、岗位职责说明',
        '四、SOP工作流程',
        '五、附录：组织架构图说明',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Number')
        set_run_font(p.runs[0], size=12)

    doc.add_page_break()

def add_position_code_table(doc):
    """添加标准岗位代码对照表"""
    add_heading_custom(doc, '一、标准岗位代码对照表', level=1)

    # 市场部
    p = doc.add_paragraph()
    set_run_font(p.add_run('1. 市场部（招商中心）— 营销推广'), size=12, bold=True)

    table1 = doc.add_table(rows=4, cols=3)
    table1.style = 'Table Grid'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['岗位代码', '岗位名称', '职责概要']
    for i, h in enumerate(headers):
        cell = table1.rows[0].cells[i]
        set_cell_shading(cell, 'E7E6E6')
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, bold=True, size=11)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    market_data = [
        ('Ta', '流量招商', '外部流量方招商对接'),
        ('La', '直播招商', '外部直播间招商对接'),
        ('Ba', '品牌招商', '外部品牌方招商对接'),
    ]
    for idx, (code, name, desc) in enumerate(market_data, 1):
        row = table1.rows[idx]
        row.cells[0].text = code
        row.cells[1].text = name
        row.cells[2].text = desc
        for cell in row.cells:
            set_run_font(cell.paragraphs[0].runs[0], size=11)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 运营部
    p = doc.add_paragraph()
    set_run_font(p.add_run('2. 运营部（运营中心）— 优化监测'), size=12, bold=True)

    table2 = doc.add_table(rows=7, cols=3)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table2.rows[0].cells[i]
        set_cell_shading(cell, 'E7E6E6')
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, bold=True, size=11)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    ops_data = [
        ('TB', '流量运营专员', '流量渠道运营对接与落地执行（合并岗位）'),
        ('Lb', '直播对接专员', '直播间运营对接与维护'),
        ('Lc', '直播运营专员', '直播运营落地执行'),
        ('Bb', '品牌对接专员', '品牌商家运营对接与维护'),
        ('Bc', '品牌运营专员', '品牌运营落地执行'),
    ]
    for idx, (code, name, desc) in enumerate(ops_data, 1):
        row = table2.rows[idx]
        row.cells[0].text = code
        row.cells[1].text = name
        row.cells[2].text = desc
        for cell in row.cells:
            set_run_font(cell.paragraphs[0].runs[0], size=11)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

def add_team_assignment(doc):
    """添加运营部岗位人员安排"""
    add_heading_custom(doc, '二、运营部岗位人员安排', level=1)

    p = doc.add_paragraph()
    set_run_font(p.add_run('生效日期：2026年5月19日 | 版本：V1.0'), size=10, color=(128,128,128))

    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['序号', '姓名', '岗位代码', '标准岗位名称']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, 'E7E6E6')
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, bold=True, size=11)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    team_data = [
        ('1', '赵思倩', 'Bb', '品牌对接专员'),
        ('2', '李芹芹', 'Lb', '直播对接专员'),
        ('3', '林勇斌', 'Lc', '直播运营专员'),
        ('4', '刘晓曼', 'Bc', '品牌运营专员'),
        ('5', '韦明哲', 'TB', '流量运营专员'),
    ]
    for idx, (seq, name, code, pos) in enumerate(team_data, 1):
        row = table.rows[idx]
        row.cells[0].text = seq
        row.cells[1].text = name
        row.cells[2].text = code
        row.cells[3].text = pos
        for cell in row.cells:
            set_run_font(cell.paragraphs[0].runs[0], size=11)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 备注
    note = doc.add_paragraph()
    set_run_font(note.add_run('备注：'), size=10.5, bold=True)
    set_run_font(note.add_run('Tc与TB合并为同一岗位，由赵思倩一人负责；其他岗位Lb/Lc/Bb/Bc各一人负责，不重复不交叉。'), size=10.5)

    doc.add_page_break()

def add_position_descriptions(doc):
    """添加岗位职责说明"""
    add_heading_custom(doc, '三、岗位职责说明', level=1)

    positions = [
        {
            'title': 'TB · 流量运营专员（合并岗位）',
            'desc': '负责外部流量方的日常运营对接与落地执行，由赵思倩一人负责，确保流量渠道稳定运行。',
            'duties': [
                '对接外部流量方，维护合作关系',
                '执行流量投放计划与活动策划',
                '监控流量数据，分析渠道效果',
                '数据录入、整理与基础分析',
                '协调内部资源，解决流量方问题',
                '定期输出流量运营报告',
            ]
        },
        {
            'title': 'Lb · 直播对接专员',
            'desc': '负责外部直播间的运营对接与维护，提升直播转化效果。',
            'duties': [
                '对接外部直播间商家，维护合作关系',
                '监控直播数据，分析转化效果',
                '协调直播资源，优化直播流程',
                '定期输出直播运营报告',
            ]
        },
        {
            'title': 'Lc · 直播运营专员',
            'desc': '负责直播运营策略的具体落地执行，配合Lb完成日常运营工作。',
            'duties': [
                '执行直播活动策划与内容排期',
                '直播数据录入、整理与基础分析',
                '协助Lb完成日常对接事务',
                '跟踪执行效果并反馈优化建议',
            ]
        },
        {
            'title': 'Bb · 品牌对接专员',
            'desc': '负责外部品牌商家的运营对接与维护，推动品牌合作落地。',
            'duties': [
                '对接外部品牌方，维护合作关系',
                '监控品牌运营数据，分析合作效果',
                '协调品牌资源，优化合作流程',
                '定期输出品牌运营报告',
            ]
        },
        {
            'title': 'Bc · 品牌运营专员',
            'desc': '负责品牌运营策略的具体落地执行，配合Bb完成日常运营工作。',
            'duties': [
                '执行品牌推广计划与活动策划',
                '品牌数据录入、整理与基础分析',
                '协助Bb完成日常对接事务',
                '跟踪执行效果并反馈优化建议',
            ]
        },
    ]

    for pos in positions:
        add_heading_custom(doc, pos['title'], level=2)
        p = doc.add_paragraph()
        set_run_font(p.add_run(pos['desc']), size=10.5)

        for duty in pos['duties']:
            bullet = doc.add_paragraph(style='List Bullet')
            set_run_font(bullet.add_run(duty), size=10.5)

        doc.add_paragraph()

def add_sop_workflow(doc):
    """添加SOP工作流程"""
    add_heading_custom(doc, '四、SOP工作流程', level=1)

    workflows = [
        {
            'title': '4.1 日常对接流程',
            'steps': [
                '每日晨会：各岗位汇报昨日工作进展及今日计划',
                '对接处理：按岗位分工处理外部合作方事务',
                '数据记录：及时录入运营数据至指定系统',
                '问题上报：遇到无法解决的问题，逐级上报',
                '日报提交：每日下班前提交工作日报',
            ]
        },
        {
            'title': '4.2 数据运营流程',
            'steps': [
                '数据采集：按规范采集各渠道/直播间/品牌数据',
                '数据清洗：整理并校验数据准确性',
                '数据分析：运用工具进行基础数据分析',
                '报告输出：按模板输出日/周/月报',
                '归档存档：报告归档至指定文件夹',
            ]
        },
        {
            'title': '4.3 跨部门协作流程',
            'steps': [
                '需求确认：与市场部确认招商需求及对接事项',
                '任务分配：运营部负责人分配具体执行任务',
                '执行跟踪：执行岗按计划推进，运营岗跟踪进度',
                '结果反馈：执行完成后向运营岗反馈结果',
                '复盘优化：定期复盘协作流程，持续优化',
            ]
        },
    ]

    for wf in workflows:
        add_heading_custom(doc, wf['title'], level=2)
        for idx, step in enumerate(wf['steps'], 1):
            p = doc.add_paragraph()
            set_run_font(p.add_run(f'{idx}. '), size=10.5, bold=True)
            set_run_font(p.add_run(step), size=10.5)

def add_appendix(doc):
    """添加附录"""
    doc.add_page_break()
    add_heading_custom(doc, '五、附录：组织架构图说明', level=1)

    p = doc.add_paragraph()
    set_run_font(p.add_run('组织架构层级说明：'), size=11, bold=True)

    items = [
        '市场部（招商中心）负责外部合作方的招商引入，对应岗位代码为Ta/La/Ba；',
        '运营部（运营中心）负责已引入合作方的日常运营维护，对应岗位代码为Tb/Tc/Lb/Lc/Bb/Bc；',
        '字母T代表流量业务线，L代表直播业务线，B代表品牌业务线；',
        '字母a代表招商岗，b代表运营岗，c代表执行岗；',
        '运营部与市场部之间通过箭头双向协作，确保招商与运营无缝衔接。',
    ]
    for item in items:
        bullet = doc.add_paragraph(style='List Bullet')
        set_run_font(bullet.add_run(item), size=10.5)

def main():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(10.5)

    # 设置页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    add_cover(doc)
    add_toc_placeholder(doc)
    add_position_code_table(doc)
    add_team_assignment(doc)
    add_position_descriptions(doc)
    add_sop_workflow(doc)
    add_appendix(doc)

    output_path = 'D:\\全犀平台\\运营管理\\运营部岗位安排SOP.docx'
    doc.save(output_path)
    print(f'SOP文档已生成：{output_path}')

if __name__ == '__main__':
    main()
